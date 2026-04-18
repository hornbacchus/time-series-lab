#!/usr/bin/env python3
"""
generate_user_guide.py - Build Word User Guide (.docx) from technique + UDF registries.

Reads:
  - resources/catalog/techniques_catalog.json
  - resources/catalog/udf_catalog.json
  - resources/techniques_md/*.md

Outputs:
  - docs/UserGuide_Source.md (generated markdown)
  - docs/TimeSeriesLab_UserGuide.docx (Word document)
"""

import json
import os
import sys
from pathlib import Path

def find_repo_root():
    """Find repo root by looking for TimeSeriesLab.sln."""
    p = Path(__file__).resolve().parent.parent
    if (p / "TimeSeriesLab.sln").exists():
        return p
    p = Path.cwd()
    while p != p.parent:
        if (p / "TimeSeriesLab.sln").exists():
            return p
        p = p.parent
    return Path(__file__).resolve().parent.parent

REPO = find_repo_root()
CATALOG_PATH = REPO / "resources" / "catalog" / "techniques_catalog.json"
UDF_CATALOG_PATH = REPO / "resources" / "catalog" / "udf_catalog.json"
TECHNIQUES_MD_DIR = REPO / "resources" / "techniques_md"
OUTPUT_MD = REPO / "docs" / "UserGuide_Source.md"
OUTPUT_DOCX = REPO / "docs" / "TimeSeriesLab_UserGuide.docx"
OUTPUT_HTML = REPO / "docs" / "TimeSeriesLab_UserGuide.html"


def load_catalog():
    if CATALOG_PATH.exists():
        with open(CATALOG_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"techniques": []}


def load_udf_catalog():
    if UDF_CATALOG_PATH.exists():
        with open(UDF_CATALOG_PATH, "r", encoding="utf-8-sig") as f:
            return json.load(f)
    return {"udfs": []}


def load_technique_md(technique_id):
    md_path = TECHNIQUES_MD_DIR / f"{technique_id}.md"
    if md_path.exists():
        return md_path.read_text(encoding="utf-8")
    return "(No detailed description available.)"


def _embed_csv_tail(lines, csv_path, n):
    """Embed the last n rows of a CSV as a markdown table."""
    import csv as csvmod
    if not csv_path.exists():
        lines.append("*(CSV file not found)*\n")
        return
    with open(csv_path, "r", encoding="utf-8") as f:
        rows = list(csvmod.reader(f))
    if len(rows) < 2:
        return
    header = rows[0]
    data = rows[-n:]
    lines.append("| " + " | ".join(header) + " |\n")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |\n")
    for row in data:
        lines.append("| " + " | ".join(row) + " |\n")


def _embed_csv_full(lines, csv_path):
    """Embed an entire CSV as a markdown table."""
    import csv as csvmod
    if not csv_path.exists():
        lines.append("*(CSV file not found)*\n")
        return
    with open(csv_path, "r", encoding="utf-8") as f:
        rows = list(csvmod.reader(f))
    if len(rows) < 2:
        return
    header = rows[0]
    lines.append("| " + " | ".join(header) + " |\n")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |\n")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |\n")


def generate_markdown(catalog, udf_catalog):
    """Generate the full UserGuide_Source.md."""
    lines = []

    def h1(text): lines.append(f"# {text}\n")
    def h2(text): lines.append(f"## {text}\n")
    def h3(text): lines.append(f"### {text}\n")
    def p(text): lines.append(f"{text}\n")
    def blank(): lines.append("")

    h1("Time Series Lab - User Guide")
    blank()
    p("Version 1.0 | Generated automatically from technique and UDF registries.")
    blank()

    # Table of Contents
    h2("Table of Contents")
    p("1. Quick Start")
    p("2. Ribbon Reference")
    p("3. How Selection Works")
    p("4. Time Index Detection")
    p("5. Frequency and Resampling")
    p("6. Presets: Fast / Balanced / Thorough")
    p("7. AUTO vs THOROUGH Formulas")
    p("8. Common Workflows")
    p("9. Audit Log and Reproducibility")
    p("10. Technique Catalog")
    p("11. UDF Reference")
    p("12. Sample Data")
    p("13. External Dependencies")
    p("14. Troubleshooting")
    blank()

    # Quick Start
    h2("1. Quick Start (3 Minutes)")
    p("1. Open Excel. The **Time Series Lab** tab appears in the ribbon.")
    p("2. Select your data columns (dates in the leftmost column, values in adjacent columns).")
    p("3. Click a **Quick Action** button (e.g., Seasonal Adjustment, Forecast).")
    p("4. The Task Pane opens with your data pre-loaded. Click **Run**.")
    p("5. Results appear in a new worksheet, along with an Audit sheet for governance.")
    blank()
    p("**Tip:** You can also use worksheet formulas (UDFs) like `=TSL_FORECAST(A2:A100, B2:B100, 12)` directly in cells.")
    blank()

    # Ribbon Reference
    h2("2. Ribbon Reference")
    p("The **Time Series Lab** ribbon tab is organized into four groups:")
    blank()
    h3("Quick Actions")
    p("One-click access to the most common analyses:")
    p("- **Seasonal Adjustment** - Decompose a series into trend, seasonal, and remainder (STL)")
    p("- **Granger Causality** - Test whether one series helps predict another")
    p("- **Lead-Lag Finder** - Find the optimal time delay between two series")
    p("- **Forecast** - Generate forecasts with prediction intervals (Auto ARIMA)")
    p("- **Anomaly Scan** - Detect outliers and anomalous observations")
    blank()
    h3("Explore")
    p("Three ways to discover and choose techniques:")
    p("- **Technique Explorer** (split button) - Browse all techniques. "
      "Click the dropdown arrow to jump directly to one of the 13 technique categories: "
      "Decomposition, Forecasting, Stationarity, Multivariate, State Space, Regimes, "
      "Volatility, Frequency Domain, Change Points, Causality, Evaluation, Missing Data, or ML/Deep Learning.")
    p("- **Recommender** (dropdown menu) - Choose an analysis goal from the dropdown "
      "(Forecast, Describe/Decompose, Detect Anomalies, Test Causality, Explore Relationships, "
      "or Statistical Testing) to start the recommendation wizard with that goal pre-selected. "
      "Or select 'Open full wizard' to walk through all 7 questions.")
    p("- **Data Readiness** - Score your selected data for quality issues before analysis.")
    blank()
    h3("Run")
    p("- **Preset** dropdown - Switch between Fast, Balanced (default), and Thorough presets")
    p("- **Run** - Execute the current technique from the Task Pane")
    p("- **Cancel** - Stop a running computation immediately")
    p("- **Re-run Thorough** - Recompute all THOROUGH formulas in the workbook")
    p("- **Settings** - Open the settings panel (presets, seed, fill method, etc.)")
    blank()
    h3("Help")
    p("- **UDF Formula Guide** - Browse all worksheet functions with examples and a formula builder")
    p("- **User Guide** - Open this document")
    p("- **About** - Version information, engine status, and diagnostics")
    blank()

    # Selection
    h2("3. How Selection Works")
    p("Time Series Lab supports **non-adjacent selections** (hold Ctrl and click multiple columns).")
    p("**Default behavior:** Each selected column is treated as one time series.")
    p("**Row mode:** Available as an override in the Task Pane for panel/cross-sectional data.")
    blank()
    p("**Headers:** The cell directly above your first selected row is used as the series name. "
      "If it looks numeric or is missing, a fallback name like `Col_B` is used.")
    blank()
    p("**Numeric coercion:** By default, string values that look like numbers are coerced. "
      "Non-numeric values become NA. This is configurable in Settings.")
    blank()
    h3("Selecting Non-Adjacent Columns (Multi-Series Techniques)")
    p("Many techniques require two or more data columns that may not be next to each other. "
      "For example, PCA, VAR, VECM, Johansen Cointegration, and Dynamic Factor Model all "
      "need multiple series as input. Time Series Lab fully supports non-adjacent column selection:")
    blank()
    p("**From the Task Pane (recommended):**")
    p("1. Click the first data column (e.g., B2:B100).")
    p("2. Hold **Ctrl** and click additional columns (e.g., D2:D100, F2:F100).")
    p("3. The selection status bar in the Task Pane will show \"3 series, 99 pts\".")
    p("4. Choose a multivariate technique (e.g., PCA, VAR) and click **Run**.")
    p("5. All selected columns are sent to the engine as separate series.")
    blank()
    p("**From a THOROUGH formula:**")
    p("The `TSL_RUN_THR` function accepts up to 5 additional data range parameters "
      "(`data_range_2` through `data_range_5`) after the `options_json` parameter:")
    blank()
    p("```")
    p("=TSL_RUN_THR(\"pca\", A2:A100, B2:B100, TSL_TRIGGER(), \"{}\", D2:D100, F2:F100, H2:H100)")
    p("```")
    blank()
    p("In this example, columns B, D, F, and H are passed as four separate series to the "
      "PCA technique. The primary `data_range` (B2:B100) can also be multi-column if "
      "some columns are adjacent.")
    blank()
    p("**Techniques that benefit from multi-series input:**")
    p("- PCA, Dynamic Factor Model (3+ series recommended)")
    p("- VAR, VECM, Bayesian VAR, Johansen Cointegration (2+ series)")
    p("- Granger Causality, Cross-Correlation, DTW, Wavelet Coherence (exactly 2 series)")
    p("- ARIMAX/SARIMAX, Transfer Function (1 target + exogenous series)")
    p("- Gradient Boosting, XGBoost, Random Forest, LSTM (optional exogenous features)")
    p("- Forecast Reconciliation (hierarchical multi-series)")
    blank()

    # Time Index
    h2("4. Time Index Detection")
    p("Time Series Lab automatically detects the date/time column by scoring candidates on:")
    p("- **Parseability:** What fraction of values are valid dates?")
    p("- **Monotonicity:** Are dates in increasing order?")
    p("- **Uniqueness:** Are there duplicate dates?")
    p("- **Regularity:** Are intervals consistent (with weekend-gap handling)?")
    blank()
    p("The best candidate above a confidence threshold is selected automatically. "
      "If no candidate scores high enough, you'll be prompted to select the date column.")
    blank()
    p("**Date-only policy:** Timestamps are floored to dates. This is recorded in the audit log.")
    p("**Duplicates:** If duplicate dates exist, they are aggregated using the series aggregation rule.")
    blank()

    # Frequency
    h2("5. Frequency and Resampling")
    p("After detecting the time index, Time Series Lab suggests a frequency:")
    p("- **Business Daily:** Weekends are non-existent (not created as timestamps)")
    p("- **Calendar Daily:** Every calendar day")
    p("- **Weekly:** Uses ISO week numbering")
    p("- **Monthly, Quarterly, Annual**")
    blank()
    p("You confirm the frequency before running. If irregular/missing dates exist, "
      "confirmation is required.")
    blank()
    h3("Resampling Aggregation (per series)")
    p("- **Flow variables** (e.g., sales): SUM")
    p("- **Rate variables** (e.g., prices): MEAN")
    p("- **Stock variables** (e.g., inventory): LAST")
    blank()
    h3("Missing Data Filling")
    p("Default: **Kalman smoothing** after resampling.")
    p("- Filled points are always flagged in the output.")
    p("- If missingness > 15% or max gap > 10 periods: strong warning + user confirmation required.")
    blank()

    # Presets
    h2("6. Presets: Fast / Balanced / Thorough")
    p("| Preset | Speed | Diagnostics | Model Search | CV Folds | Use Case |")
    p("|--------|-------|-------------|-------------|----------|----------|")
    p("| **Fast** | Fastest | Minimal | Safe defaults | None | Interactive exploration |")
    p("| **Balanced** | Moderate | Standard | Moderate candidates | 3-fold | General use (default) |")
    p("| **Thorough** | Slowest | Full | Wide search | 5-fold | Publication / audit |")
    blank()
    p("**Thorough** is always manual (handle-based) in formulas. It never surprises you with long recalculation.")
    blank()

    # AUTO vs THOROUGH
    h2("7. AUTO vs THOROUGH Formulas")
    p("Time Series Lab provides two lanes of worksheet formulas:")
    blank()
    p("| Feature | AUTO | THOROUGH |")
    p("|---------|------|----------|")
    p("| Function prefix | `TSL_*` | `TSL_*_THR` or `TSL_RUN_THR` |")
    p("| Category | \"AUTO (recomputes)\" | \"THOROUGH (manual handles)\" |")
    p("| Description | Starts with \"AUTO:\" | Starts with \"THOROUGH:\" |")
    p("| Presets | Fast or Balanced only | Thorough only |")
    p("| Output | Numbers/arrays (spill) | Handle string (TSL.THR.*) |")
    p("| Recalculation | Normal Excel recalc | Only when trigger changes |")
    p("| Trigger | Not needed | REQUIRED: `TSL_TRIGGER()` |")
    blank()
    h3("Worked Example: AUTO Forecast")
    p("```")
    p("=TSL_FORECAST(A2:A100, B2:B100, 12, \"Balanced\", \"Auto\")")
    p("```")
    p("This spills 12 rows of forecasts directly. It recomputes whenever inputs change.")
    blank()
    h3("Worked Example: THOROUGH Forecast")
    p("```")
    p("Cell D1: =TSL_FORECAST_THR(A2:A100, B2:B100, 12, TSL_TRIGGER(), \"Auto\")")
    p("Cell E1: =TSL_STATUS(D1)")
    p("Cell F1: =TSL_WARNINGS(D1)")
    p("```")
    p("D1 returns a handle like `TSL.THR.auto_arima.20260216143022.a1b2c3d4`.")
    p("Use `=TSL_TABLE(D1, \"forecasts\")` to extract the forecast table.")
    p("Click **Re-run Thorough** in the ribbon to recompute.")
    blank()

    # Common Workflows
    h2("8. Common Workflows")
    h3("Seasonal Adjustment")
    p("1. Select date + value columns")
    p("2. Click **Seasonal Adjustment** in Quick Actions")
    p("3. Choose method (STL recommended), confirm frequency")
    p("4. Click Run. Review trend, seasonal, and remainder components.")
    blank()
    h3("Granger Causality")
    p("1. Select two series (potential cause and effect)")
    p("2. Click **Granger Causality**")
    p("3. Set max lags (default 12)")
    p("4. Review F-statistics and p-values by lag")
    blank()
    h3("Lead-Lag Finder")
    p("1. Select two series")
    p("2. Click **Lead-Lag Finder**")
    p("3. Choose method (Prewhitened CCF recommended)")
    p("4. Review best lag and correlation strength")
    blank()
    h3("Forecasting")
    p("1. Select date + value column")
    p("2. Click **Forecast**")
    p("3. Set horizon, choose model (Auto recommended)")
    p("4. Review forecasts with prediction intervals")
    blank()

    # Audit
    h2("9. Audit Log and Reproducibility")
    p("Every run creates:")
    p("- **Results sheet:** Plain-English summary, diagnostics, output tables, charting suggestions")
    p("- **Audit sheet:** Full record of inputs, parameters, transforms, diagnostics, versions, seed")
    p("- **Embedded JSON:** Machine-readable record in hidden sheet `_TSL_RUNS`")
    blank()
    p("The JSON record contains everything needed to reproduce a run: input ranges, resolved "
      "parameters, seed, and version information. Use **Re-run this analysis** from the Task Pane.")
    blank()

    # Technique Catalog
    h2("10. Technique Catalog")
    techniques = catalog.get("techniques", [])
    categories = {}
    for t in techniques:
        cat = t.get("category", "Other")
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(t)

    for cat in sorted(categories.keys()):
        h3(cat)
        for t in categories[cat]:
            tid = t.get("id", "?")
            name = t.get("name", tid)
            summary = t.get("summary", "")
            lines.append(f"#### {name}\n")
            blank()
            p(f"`{tid}`: {summary}")
            blank()

            # Load markdown description — downgrade headings so they nest
            # under the technique H4 (# -> #####, ## -> #####)
            desc = load_technique_md(tid)
            if desc and desc != "(No detailed description available.)":
                for dl in desc.split("\n"):
                    if dl.startswith("## "):
                        lines.append(f"##### {dl[3:]}\n")
                    elif dl.startswith("# "):
                        lines.append(f"##### {dl[2:]}\n")
                    else:
                        lines.append(f"{dl}\n")
                blank()

            # Parameters
            params = t.get("parameters", [])
            if params:
                lines.append("##### Parameters\n")
                for param in params:
                    pname = param.get("name", "?")
                    plabel = param.get("label", pname)
                    pdef = param.get("default", "")
                    pdesc = param.get("description", "")
                    adv = " *(advanced)*" if param.get("advanced", False) else ""
                    p(f"- `{pname}` ({plabel}): {pdesc} Default: {pdef}{adv}")
                blank()

    # UDF Reference
    h2("11. UDF Reference")
    udfs = udf_catalog.get("udfs", [])
    if udfs:
        for udf in udfs:
            name = udf.get("name", "?")
            desc = udf.get("description", "")
            lane = udf.get("lane", "?")
            h3(f"`{name}` [{lane}]")
            p(desc)
            blank()
            args = udf.get("arguments", [])
            if args:
                p("Arguments:")
                for arg in args:
                    aname = arg.get("name", "?")
                    adesc = arg.get("description", "")
                    p(f"- `{aname}`: {adesc}")
                blank()
    else:
        p("*(UDF catalog not yet generated. Run `generate_udf_catalog.ps1` first.)*")
    blank()

    # Sample Data
    h2("12. Sample Data")
    p("The following datasets are included with Time Series Lab for testing and exploration. "
      "Each is provided as a CSV file in `resources/sample_data/` and can be copy-pasted "
      "directly into Excel.")
    blank()

    h3("Treasury Constant Maturity Yields (Daily)")
    p("Daily 2-year, 5-year, 10-year, and 30-year U.S. Treasury constant maturity yields "
      "from the Federal Reserve (FRED series DGS2, DGS5, DGS10, DGS30). "
      "Date range depends on tenor: 10Y and 5Y begin January 1962; 2Y begins June 1976; "
      "30Y begins February 1977.")
    blank()
    p("**File:** `resources/sample_data/treasury_yields.csv`")
    p("**Rows:** ~16,000 (business days)")
    p("**Suggested techniques:** Forecasting (Auto ARIMA, ETS), Cointegration (Johansen, VECM), "
      "PCA, Granger Causality, Lead-Lag, Volatility (GARCH)")
    blank()

    # Embed a small preview of the Treasury data
    p("**Preview (last 10 rows):**")
    blank()
    _embed_csv_tail(lines, REPO / "resources" / "sample_data" / "treasury_yields.csv", 10)
    blank()

    h3("U.S. Real GDP Growth (Quarterly, SAAR)")
    p("Quarter-over-quarter annualized growth rate of real U.S. GDP from the Bureau of "
      "Economic Analysis (FRED series A191RL1Q225SBEA). Begins Q1 1950.")
    blank()
    p("**File:** `resources/sample_data/real_gdp.csv`")
    p("**Suggested techniques:** Forecasting, Regime Switching (Markov, SETAR), "
      "Change Point Detection, Decomposition (STL)")
    blank()
    p("**Full dataset:**")
    blank()
    _embed_csv_full(lines, REPO / "resources" / "sample_data" / "real_gdp.csv")
    blank()

    h3("Core PCE Inflation (Quarterly, SAAR)")
    p("Quarter-over-quarter annualized percent change in the core Personal Consumption "
      "Expenditures price index (excluding food and energy) from the Bureau of Economic "
      "Analysis (FRED series JCXFE). Begins Q1 1959.")
    blank()
    p("**File:** `resources/sample_data/core_pce.csv`")
    p("**Suggested techniques:** Forecasting, Granger Causality (with GDP or yields), "
      "Structural Breaks (Zivot-Andrews), Stationarity Tests (ADF, KPSS)")
    blank()
    p("**Full dataset:**")
    blank()
    _embed_csv_full(lines, REPO / "resources" / "sample_data" / "core_pce.csv")
    blank()

    h3("Total Nonfarm Payrolls - Seasonally Adjusted (Monthly)")
    p("U.S. total nonfarm payroll employment, seasonally adjusted levels, in thousands "
      "of jobs. From the Bureau of Labor Statistics Current Employment Statistics (CES) "
      "program (FRED series PAYEMS). Monthly, January 1939 to present.")
    blank()
    p("**File:** `resources/sample_data/nonfarm_payroll_sa.csv`")
    p("**Rows:** ~1,047 months")
    p("**Suggested techniques:** Forecasting (Auto ARIMA, ETS), Change Point Detection, "
      "Regime Switching, Structural Breaks")
    blank()
    p("**Preview (last 10 rows):**")
    blank()
    _embed_csv_tail(lines, REPO / "resources" / "sample_data" / "nonfarm_payroll_sa.csv", 10)
    blank()

    h3("Total Nonfarm Payrolls - Not Seasonally Adjusted (Monthly)")
    p("Same series as above, but without seasonal adjustment (FRED series PAYNSA). This "
      "is the ideal input for seasonal-adjustment techniques (X-13 ARIMA-SEATS, STL): "
      "run the decomposition on this NSA series and compare the resulting SA column "
      "against the official PAYEMS series above.")
    blank()
    p("**File:** `resources/sample_data/nonfarm_payroll_nsa.csv`")
    p("**Rows:** ~1,047 months")
    p("**Suggested techniques:** Seasonal Adjustment (X-13 ARIMA-SEATS, STL, Classical), "
      "Decomposition, Frequency Domain (spectral peaks at period 12)")
    blank()
    p("**Note:** This is a *levels* series (total employment in thousands), not a "
      "month-over-month change series. Apply seasonal adjustment to levels first, "
      "then compute job gains as first differences of the SA column if needed.")
    blank()
    p("**Preview (last 10 rows):**")
    blank()
    _embed_csv_tail(lines, REPO / "resources" / "sample_data" / "nonfarm_payroll_nsa.csv", 10)
    blank()

    # External Dependencies
    h2("13. External Dependencies")
    p("Time Series Lab bundles its own Python runtime and most required libraries. "
      "However, one external binary and four optional Python packages may need to be "
      "installed separately depending on which techniques you plan to use. The "
      "**Help -> About** dialog reports what is currently detected on your machine.")
    blank()

    h3("X-13 ARIMA-SEATS Binary (required for Seasonal Adjustment)")
    p("X-13 ARIMA-SEATS is the U.S. Census Bureau's seasonal adjustment program. "
      "It is the statistical backend for the **Seasonal Adjustment** Quick Action when "
      "the technique **X-13 ARIMA-SEATS** is selected. Because the binary has its own "
      "license, it is not bundled with the installer - you must download it separately "
      "from Census and drop it into the project's `resources/x13/` folder.")
    blank()
    p("**Download:** https://www.census.gov/data/software/x13as.html")
    p("**Direct archive:** "
      "https://www2.census.gov/software/x-13arima-seats/x13as/windows/program-archives/")
    blank()
    p("**Installation (Windows):**")
    p("1. Download `x13as_html-v1-1-b62.zip` (HTML-output build) or the ASCII-output "
      "equivalent.")
    p("2. Extract the archive. You will get a folder named `x13as/` containing "
      "`x13as_html.exe` (or `x13as_ascii.exe`).")
    p("3. Copy the `.exe` directly into `resources/x13/` under your Time Series Lab "
      "project root (the `x13as/` wrapper folder is not needed - the engine searches "
      "for the executable one level down).")
    p("4. Restart Excel. The **Help -> About** dialog will confirm that the X-13 binary "
      "is detected.")
    blank()
    p("**Accepted filenames** (the engine searches for any of these, in order): "
      "`x13as_html.exe`, `x13as_ascii.exe`, `x13ashtml.exe`, `x13as.exe`. On macOS "
      "and Linux, drop the filename extension and make the file executable with "
      "`chmod +x`.")
    blank()
    p("**Without the X-13 binary:** the X-13 ARIMA-SEATS technique returns a clear "
      "error with download instructions. All other seasonal decomposition techniques "
      "(STL, Classical, MSTL) work without any external binary and are reasonable "
      "substitutes for exploratory work.")
    blank()

    h3("Core Python Packages (bundled, always installed)")
    p("These ten packages ship with the installer and cover 58 of the 67 techniques. "
      "You should never need to install them manually:")
    blank()
    p("| Package | Used by |")
    p("| --- | --- |")
    p("| `numpy`, `scipy`, `pandas` | Core numeric and data structures |")
    p("| `statsmodels` | ARIMA, VAR, VECM, GARCH, X-13 wrapper, unit root tests |")
    p("| `pmdarima` | Auto ARIMA |")
    p("| `arch` | GARCH, EGARCH, GJR-GARCH volatility models |")
    p("| `ruptures` | Change point detection (Pelt, BinSeg, Window) |")
    p("| `PyWavelets` | Wavelet decomposition, wavelet coherence |")
    p("| `hmmlearn` | Hidden Markov regime-switching models |")
    p("| `scikit-learn` | Isolation Forest, PCA, clustering, MLPRegressor |")
    blank()

    h3("Optional Python Packages (install for full ML/DL fidelity)")
    p("Nine of the 67 techniques use deep learning or specialized ML libraries. "
      "These are marked optional - if not installed, each technique **gracefully "
      "falls back** to a simpler backend and returns a warning explaining what "
      "was substituted, so the tool always produces a result.")
    blank()
    p("| Package | Techniques | Fallback when missing |")
    p("| --- | --- | --- |")
    p("| `prophet` | `prophet_forecast` | Seasonal naive |")
    p("| `xgboost` | `xgboost_forecast` | `sklearn.GradientBoostingRegressor` |")
    p("| `reservoirpy` | `echo_state_network` | Minimal numpy reservoir |")
    p("| `torch` (CPU) | `lstm_gru_forecast`, `nbeats_forecast`, `nhits_forecast`, "
      "`tcn_forecast`, `transformer_forecast`, `autoencoder_anomaly` | `sklearn.MLPRegressor` "
      "or `IsolationForest` |")
    blank()
    p("**To install all four** (open a Command Prompt or PowerShell):")
    blank()
    p("```")
    p("pip install prophet xgboost reservoirpy")
    p("pip install torch --index-url https://download.pytorch.org/whl/cpu")
    p("```")
    blank()
    p("Notes:")
    p("- **`torch`** is the largest (~200 MB for the CPU-only wheel, ~2.5 GB with CUDA). "
      "Use the CPU-only index URL above unless you have a CUDA GPU you want to use.")
    p("- **`prophet`** installs `cmdstanpy` and compiles a small C++ Stan backend on "
      "first import (~80 MB). A one-time cost.")
    p("- **`reservoirpy`** and **`xgboost`** are small (~50 MB each) and install without "
      "compilation.")
    p("- The installer offers a post-install **Install Optional ML/DL Packages** check "
      "box that runs these commands for you.")
    blank()

    h3("Multiple Python Installs on Your Machine")
    p("If your PC has more than one Python installation (common on Windows when "
      "Python is installed from python.org AND the Microsoft Store AND via a "
      "standalone runtime AND via Anaconda), `pip install <package>` may install "
      "the package into a Python different from the one Time Series Lab actually "
      "uses. When you then run a technique, the engine reports something like "
      "`No module named 'statsmodels'` even though `pip list` swears the package "
      "is installed.")
    blank()
    p("The engine launches Python via Windows PATH resolution (the first `python.exe` "
      "found). To install packages into that specific interpreter, do not rely on "
      "a bare `pip install ...`. Instead, call pip through the exact python.exe "
      "the engine will use:")
    blank()
    p("```")
    p("# Find which Python the engine will launch")
    p("where python")
    p("")
    p("# Install into THAT Python specifically (adjust path to match)")
    p("\"C:\\Python314\\python.exe\" -m pip install -r ^")
    p("    \"<project>\\engine\\requirements.txt\"")
    p("\"C:\\Python314\\python.exe\" -m pip install prophet xgboost reservoirpy")
    p("\"C:\\Python314\\python.exe\" -m pip install torch ^")
    p("    --index-url https://download.pytorch.org/whl/cpu")
    p("```")
    blank()
    p("If the installer provisioned a bundled runtime at "
      "`%LOCALAPPDATA%\\TimeSeriesLab\\engine\\runtime\\python.exe`, the engine "
      "uses that first and this entire issue disappears — pip-installing into the "
      "bundled runtime updates exactly the Python the engine sees.")
    blank()

    h3("What Does NOT Need To Be Installed Separately")
    p("Time Series Lab does not require R, TRAMO/SEATS, JDemetra+, EViews, SAS, MATLAB, "
      "or any cloud service. Everything runs locally. The only external components "
      "are the X-13 binary (optional, for one technique) and the four optional Python "
      "packages listed above.")
    blank()

    h3("Checking What Is Installed")
    p("**Help -> About** in the ribbon reports:")
    p("- Add-in version and engine version")
    p("- Python runtime version")
    p("- Engine process status (running / not running)")
    p("- Technique library size (techniques x categories)")
    p("- Engine pipe status")
    blank()
    p("Planned for a future release: a **Dependency Doctor** panel that inventories "
      "every required and optional package, flags what is missing, and generates the "
      "exact install commands for your system.")
    blank()

    # Troubleshooting
    h2("14. Troubleshooting")
    p("Time Series Lab follows a **fail loudly** policy. Errors include plain-English "
      "explanations and suggested fixes.")
    blank()
    p("**Common issues:**")
    p("- **\"No range selected\"**: Select data columns before clicking Run.")
    p("- **\"Could not detect time index\"**: Ensure dates are in a column adjacent to your data.")
    p("- **\"Engine not found\"**: Reinstall the add-in or check the Python runtime at "
      "`%LOCALAPPDATA%\\TimeSeriesLab\\engine\\runtime\\`.")
    p("- **\"Thorough preset not allowed in AUTO\"**: Use `TSL_RUN_THR()` or the Task Pane for Thorough.")
    p("- **\"MISSING (Click Re-run Thorough)\"**: The handle result is not in cache. "
      "Click Re-run Thorough in the ribbon.")
    p("- **Missing data warning**: If >15% missing or gap >10 periods, confirm before proceeding.")
    p("- **\"X-13 ARIMA-SEATS binary not found\"**: Download the Census X-13 executable "
      "and place it in `resources/x13/`. See **Section 13: External Dependencies** for "
      "step-by-step instructions.")
    p("- **X-13 returned \"Number of years spanned exceeds program limit (85)\"**: Your "
      "series is longer than 85 years, which is X-13's hard cap. The engine truncates "
      "automatically to satisfy the limit; the warning in the result tells you how many "
      "observations were dropped. You can also set `fit_window_obs` in the Run panel "
      "to use a shorter window (BLS-style concurrent adjustment uses ~120 months).")
    p("- **X-13 \"Estimation failed to converge\"**: Automatic model selection could not "
      "fit your series. The engine automatically retries with progressively simpler "
      "ARIMA specifications (airline, MA-only, seasonal-MA-only) and reports which one "
      "succeeded in the warnings. If all fail, the series is likely pre-differenced or "
      "too structurally variable - try a shorter `fit_window_obs` or use STL Decomposition "
      "as an alternative.")
    p("- **Optional ML technique says \"Falling back to ...\"**: Install the matching "
      "optional package (see **Section 13: External Dependencies**) to use the full "
      "model backend. The fallback result is statistically valid but typically less "
      "accurate than the native backend.")
    blank()

    return "\n".join(lines)


def generate_docx(markdown_text):
    """Generate Word document from markdown text using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("WARNING: python-docx not installed. Skipping .docx generation.")
        print("  Install with: pip install python-docx")
        return False

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Process markdown line by line
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# ") and not line.startswith("## "):
            # Title
            heading = doc.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("```"):
            # Code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.style = doc.styles['Normal']
        elif line.startswith("| ") and "---" not in line:
            # Table row - collect all table rows
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                if "---" not in lines[i]:
                    table_lines.append(lines[i])
                i += 1
            i -= 1  # Will be incremented at end of loop

            if len(table_lines) > 1:
                # Parse table
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                rows = []
                for tl in table_lines[1:]:
                    cells = [c.strip() for c in tl.split("|")[1:-1]]
                    rows.append(cells)

                ncols = len(headers)
                table = doc.add_table(rows=1 + len(rows), cols=ncols)
                table.style = 'Table Grid'

                # Headers
                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = h.replace("**", "")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Data rows
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < ncols:
                            table.rows[r_idx + 1].cells[c_idx].text = cell_text.replace("**", "").replace("`", "")

        elif line.startswith("- "):
            # Bullet list
            text = line[2:].strip()
            # Remove markdown bold/code formatting
            text = text.replace("**", "").replace("`", "")
            doc.add_paragraph(text, style='List Bullet')
        elif line.strip() and not line.startswith("|"):
            # Regular paragraph
            text = line.replace("**", "").replace("`", "")
            doc.add_paragraph(text)

        i += 1

    # Save
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))
    return True


def generate_html(markdown_text):
    """Generate a self-contained HTML file with sidebar nav and OkabeIto styling."""
    import re

    # --- First pass: collect headings for sidebar nav ---
    # Sidebar shows only H2 (main sections) and H3 (technique categories).
    # Individual techniques and sub-headings stay in content only.
    sections = []  # list of (level, slug, title)
    lines = markdown_text.split("\n")
    in_catalog = False
    for line in lines:
        if line.startswith("## "):
            title = line[3:].strip()
            in_catalog = 'technique catalog' in title.lower()
            sections.append((2, _slugify(title), title))
        elif line.startswith("### ") and in_catalog:
            title = line[4:].strip()
            sections.append((3, _slugify(title), title))

    sidebar_html = []
    sidebar_html.append('<nav>')
    sidebar_html.append('  <div class="logo">')
    sidebar_html.append('    <h1>Time Series Lab</h1>')
    sidebar_html.append('    <span>User Guide v1.0</span>')
    sidebar_html.append('  </div>')
    for level, slug, title in sections:
        if level == 2:
            short = re.sub(r'^\d+\.\s*', '', title)
            sidebar_html.append(f'  <a href="#{slug}" class="section-head">{_esc(short)}</a>')
        elif level == 3:
            sidebar_html.append(f'  <a href="#{slug}">&nbsp;&nbsp;{_esc(title)}</a>')
    sidebar_html.append('</nav>')

    # --- Second pass: render content ---
    content = []
    i = 0
    in_list = False
    while i < len(lines):
        line = lines[i]

        # Close list if leaving bullet items
        if in_list and not line.startswith("- "):
            content.append("</ul>")
            in_list = False

        if line.startswith("##### "):
            title = line[6:].strip()
            content.append(f'<h5>{_md_inline(title)}</h5>')
        elif line.startswith("#### "):
            title = line[5:].strip()
            slug = _slugify(title)
            content.append(f'<h4 id="{slug}">{_md_inline(title)}</h4>')
        elif line.startswith("### "):
            title = line[4:].strip()
            slug = _slugify(title)
            content.append(f'<h3 id="{slug}">{_md_inline(title)}</h3>')
        elif line.startswith("## "):
            title = line[3:].strip()
            slug = _slugify(title)
            content.append(f'<h2 id="{slug}">{_md_inline(title)}</h2>')
        elif line.startswith("# ") and not line.startswith("## "):
            # Skip the H1 title — it's in the sidebar logo
            pass
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_html = "\n".join(_esc(cl) for cl in code_lines)
            content.append(f"<pre>{code_html}</pre>")
        elif line.startswith("| ") and "---" not in line:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("|"):
                if "---" not in lines[i]:
                    table_lines.append(lines[i])
                i += 1
            i -= 1
            if len(table_lines) > 1:
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                content.append("<table><tr>")
                for h in headers:
                    content.append(f"<th>{_md_inline(h)}</th>")
                content.append("</tr>")
                for tl in table_lines[1:]:
                    cells = [c.strip() for c in tl.split("|")[1:-1]]
                    content.append("<tr>")
                    for c in cells:
                        content.append(f"<td>{_md_inline(c)}</td>")
                    content.append("</tr>")
                content.append("</table>")
        elif line.startswith("- "):
            if not in_list:
                content.append("<ul>")
                in_list = True
            content.append(f"<li>{_md_inline(line[2:].strip())}</li>")
        elif line.strip():
            text = line.strip()
            # Detect tip/warning patterns
            if text.startswith("**Tip:**") or text.startswith("**Tip**:"):
                content.append(f'<div class="tip">{_md_inline(text)}</div>')
            elif text.startswith("**Warning:**") or text.startswith("**Warning**:"):
                content.append(f'<div class="warn">{_md_inline(text)}</div>')
            else:
                content.append(f"<p>{_md_inline(text)}</p>")

        i += 1

    if in_list:
        content.append("</ul>")

    # --- Assemble full HTML ---
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Time Series Lab - User Guide</title>
<style>
  :root {{
    --oi-blue: #0072B2;
    --oi-orange: #E69F00;
    --oi-green: #009E73;
    --oi-vermillion: #D55E00;
    --oi-skyblue: #56B4E9;
    --oi-purple: #CC79A7;
    --oi-yellow: #F0E442;
    --oi-black: #000000;
    --bg: #FAFAFA;
    --sidebar-bg: #1E293B;
    --sidebar-text: #CBD5E1;
    --sidebar-hover: #334155;
    --sidebar-active: #0072B2;
    --card-bg: #FFFFFF;
    --border: #E2E8F0;
    --text: #1E293B;
    --text-secondary: #64748B;
    --code-bg: #F1F5F9;
  }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: Verdana, Geneva, sans-serif;
    font-size: 9pt;
    background: var(--bg);
    color: var(--text);
    line-height: 1.6;
    display: flex;
    min-height: 100vh;
  }}
  nav {{
    width: 260px;
    min-width: 260px;
    background: var(--sidebar-bg);
    color: var(--sidebar-text);
    padding: 24px 0;
    position: fixed;
    top: 0;
    left: 0;
    bottom: 0;
    overflow-y: auto;
    z-index: 10;
  }}
  nav .logo {{
    padding: 0 20px 20px;
    border-bottom: 1px solid #334155;
    margin-bottom: 12px;
  }}
  nav .logo h1 {{
    font-size: 16px;
    font-weight: 700;
    color: #F8FAFC;
    letter-spacing: -0.3px;
  }}
  nav .logo span {{
    font-size: 12px;
    color: var(--oi-skyblue);
    font-weight: 400;
  }}
  nav a {{
    display: block;
    padding: 7px 20px;
    color: var(--sidebar-text);
    text-decoration: none;
    font-size: 13px;
    transition: background 0.15s, color 0.15s;
  }}
  nav a:hover {{ background: var(--sidebar-hover); color: #F8FAFC; }}
  nav a.section-head {{
    font-size: 13px;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    color: var(--oi-skyblue);
    padding-top: 18px;
    padding-bottom: 4px;
    margin-top: 4px;
    border-top: 1px solid #334155;
  }}
  nav a.section-head:hover {{
    color: #FFFFFF;
  }}
  nav a.section-head:first-of-type {{
    border-top: none;
    margin-top: 0;
  }}
  main {{
    margin-left: 260px;
    flex: 1;
    padding: 40px 48px 80px;
    max-width: 900px;
  }}
  h2 {{
    font-size: 24px;
    font-weight: 700;
    margin: 48px 0 16px;
    padding-bottom: 8px;
    border-bottom: 2px solid var(--oi-blue);
    color: var(--text);
  }}
  h2:first-of-type {{ margin-top: 0; }}
  h3 {{
    font-size: 17px;
    font-weight: 600;
    margin: 28px 0 10px;
    color: var(--text);
  }}
  h4 {{
    font-size: 15px;
    font-weight: 700;
    margin: 32px 0 8px;
    padding-bottom: 4px;
    border-bottom: 1px solid var(--border);
    color: var(--oi-blue);
  }}
  h5 {{
    font-size: 13px;
    font-weight: 600;
    margin: 18px 0 6px;
    color: var(--text-secondary);
    text-transform: uppercase;
    letter-spacing: 0.5px;
  }}
  p {{ margin-bottom: 12px; }}
  .card {{
    background: var(--card-bg);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 20px 24px;
    margin: 16px 0;
  }}
  .card h4 {{
    font-size: 14px;
    font-weight: 700;
    margin-bottom: 8px;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    margin: 12px 0 20px;
    font-size: 14px;
  }}
  th {{
    text-align: left;
    padding: 10px 12px;
    background: var(--code-bg);
    border: 1px solid var(--border);
    font-weight: 600;
    font-size: 13px;
  }}
  td {{
    padding: 8px 12px;
    border: 1px solid var(--border);
    vertical-align: top;
  }}
  code {{
    background: var(--code-bg);
    padding: 2px 6px;
    border-radius: 4px;
    font-family: 'Cascadia Code', 'Consolas', monospace;
    font-size: 13px;
  }}
  pre {{
    background: var(--code-bg);
    padding: 14px 18px;
    border-radius: 6px;
    overflow-x: auto;
    font-family: 'Cascadia Code', 'Consolas', monospace;
    font-size: 13px;
    margin: 12px 0 20px;
    line-height: 1.5;
  }}
  .tip {{
    background: #EFF6FF;
    border-left: 4px solid var(--oi-blue);
    padding: 12px 16px;
    margin: 16px 0;
    border-radius: 0 6px 6px 0;
    font-size: 14px;
  }}
  .tip strong {{ color: var(--oi-blue); }}
  .warn {{
    background: #FFFBEB;
    border-left: 4px solid var(--oi-orange);
    padding: 12px 16px;
    margin: 16px 0;
    border-radius: 0 6px 6px 0;
    font-size: 14px;
  }}
  .warn strong {{ color: var(--oi-vermillion); }}
  ul {{ margin: 8px 0 12px 20px; font-size: 14px; }}
  li {{ margin: 4px 0; }}
  footer {{
    margin-top: 60px;
    padding-top: 20px;
    border-top: 1px solid var(--border);
    color: var(--text-secondary);
    font-size: 13px;
    text-align: center;
  }}
  @media (max-width: 768px) {{
    nav {{ display: none; }}
    main {{ margin-left: 0; padding: 24px 20px; }}
  }}
</style>
</head>
<body>
{chr(10).join(sidebar_html)}
<main>
{chr(10).join(content)}
<footer>
  Time Series Lab v1.0 &bull; Created by Matthew T. Hornbach &bull; Generated automatically from technique and UDF registries
</footer>
</main>
</body>
</html>"""

    OUTPUT_HTML.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_HTML.write_text(html, encoding="utf-8")
    return True


def _slugify(text):
    """Convert heading text to a URL-safe slug."""
    import re
    text = re.sub(r'[^\w\s-]', '', text.lower())
    text = re.sub(r'[\s]+', '-', text.strip())
    return text


def _esc(text):
    """Escape HTML special characters."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _md_inline(text):
    """Convert inline markdown (bold, code) to HTML."""
    import re
    text = _esc(text)
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
    return text


def main():
    print("=== Generating Time Series Lab User Guide ===")

    catalog = load_catalog()
    udf_catalog = load_udf_catalog()

    print(f"Loaded {len(catalog.get('techniques', []))} techniques")
    print(f"Loaded {len(udf_catalog.get('udfs', []))} UDFs")

    # Generate markdown
    md = generate_markdown(catalog, udf_catalog)
    OUTPUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD.write_text(md, encoding="utf-8")
    print(f"Generated markdown: {OUTPUT_MD}")

    # Generate docx
    if generate_docx(md):
        print(f"Generated Word doc: {OUTPUT_DOCX}")
    else:
        print("Word doc generation skipped (install python-docx).")

    # Generate HTML
    if generate_html(md):
        print(f"Generated HTML: {OUTPUT_HTML}")

    print("=== Done ===")


if __name__ == "__main__":
    main()
